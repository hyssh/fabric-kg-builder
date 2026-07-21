"""DOCX extractor built on python-docx."""

from __future__ import annotations

import hashlib
import html
from datetime import datetime, timezone
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from fabric_kg_builder.model.ids import (
    content_hash as compute_content_hash,
    make_document_element_id,
    make_source_file_id,
)
from fabric_kg_builder.model.schemas import DocumentElementRow, SourceFileRow

if TYPE_CHECKING:
    from .adapter import HyperlinkRecord


class DocxExtractResult:
    """Result returned by :meth:`DocxExtractor.extract`."""

    __slots__ = ("source_file", "document_elements", "hyperlinks")

    def __init__(
        self,
        source_file: SourceFileRow,
        document_elements: list[DocumentElementRow],
        hyperlinks: "list[HyperlinkRecord] | None" = None,
    ) -> None:
        self.source_file = source_file
        self.document_elements = document_elements
        self.hyperlinks: list[HyperlinkRecord] = hyperlinks or []


def _canonical_path(path: Path, project_root: Path | None) -> str:
    root = Path(project_root) if project_root else Path.cwd()
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return path.as_posix()


def _file_content_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _extract_hyperlinks(document: DocumentObject) -> "list[HyperlinkRecord]":
    """Extract all hyperlinks from a DOCX document (EXT-003).

    Iterates every paragraph (including those inside tables) and resolves
    ``<w:hyperlink r:id="…">`` relationship IDs to their target URLs.
    Populates ``source_locator_json`` with paragraph or table-cell context.
    """
    from .adapter import HyperlinkRecord  # local import avoids circular deps
    import json as _json

    rels = document.part.rels
    links: list[HyperlinkRecord] = []

    def _collect_para_links(
        para: Paragraph,
        para_index: int,
        table_locator: "dict | None" = None,
    ) -> None:
        for hl_elem in para._p.findall(".//" + qn("w:hyperlink")):
            rid = hl_elem.get(qn("r:id"))
            target = ""
            if rid and rid in rels:
                rel = rels[rid]
                target = getattr(rel, "target_ref", "") or ""
            # Gather anchor text from all <w:t> inside the hyperlink
            texts = [
                (t.text or "")
                for t in hl_elem.findall(".//" + qn("w:t"))
            ]
            anchor = "".join(texts).strip()
            if anchor or target:
                if table_locator is not None:
                    locator = _json.dumps({"type": "docx", **table_locator})
                else:
                    locator = _json.dumps({"type": "docx", "paragraph_index": para_index})
                links.append(HyperlinkRecord(anchor=anchor, target=target, source_locator_json=locator))

    for para_index, para in enumerate(document.paragraphs):
        _collect_para_links(para, para_index=para_index)

    for tbl_idx, tbl in enumerate(document.tables):
        for row_idx, row in enumerate(tbl.rows):
            for col_idx, cell in enumerate(row.cells):
                tbl_locator = {
                    "table_index": tbl_idx,
                    "row_index": row_idx,
                    "col_index": col_idx,
                }
                for para in cell.paragraphs:
                    _collect_para_links(para, para_index=0, table_locator=tbl_locator)

    return links


def _iter_block_items(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _render_table_html(table: Table) -> str:
    rows_html: list[str] = []
    for row in table.rows:
        cell_html = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
        rows_html.append(f"<tr>{cell_html}</tr>")
    return f"<table>{''.join(rows_html)}</table>"


class DocxExtractor:
    """Extract section, paragraph, and table elements from DOCX files."""

    @staticmethod
    def extract(path: str | Path, *, project_root: str | Path | None = None) -> DocxExtractResult:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Source file not found: {path}")

        now = datetime.now(timezone.utc)
        file_hash = _file_content_hash(path)
        canonical_path = _canonical_path(path, Path(project_root) if project_root else None)
        source_file_id = make_source_file_id(canonical_path, file_hash)

        source_file = SourceFileRow(
            source_file_id=source_file_id,
            path=canonical_path,
            filename=path.name,
            source_type="docx",
            content_hash=file_hash,
            byte_size=path.stat().st_size,
            ingested_at=now,
        )

        document = Document(path)
        document_elements: list[DocumentElementRow] = []
        heading_stack: list[str] = []
        heading_id_stack: list[str] = []
        sort_order = 0

        for block in _iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue

                style_name = block.style.name if block.style is not None else ""
                if style_name.startswith("Heading "):
                    try:
                        level = int(style_name.split()[-1])
                    except ValueError:
                        level = 1
                    if level in {1, 2, 3}:
                        heading_stack = heading_stack[: level - 1]
                        heading_id_stack = heading_id_stack[: level - 1]
                        heading_stack.append(text)
                        section_path = "/".join(heading_stack)
                        section_hash = compute_content_hash(text)
                        parent_element_id = heading_id_stack[-1] if heading_id_stack else None
                        section_id = make_document_element_id(
                            source_file_id,
                            "section",
                            None,
                            sort_order,
                            section_hash,
                        )
                        document_elements.append(
                            DocumentElementRow(
                                document_element_id=section_id,
                                source_file_id=source_file_id,
                                element_type="section",
                                parent_element_id=parent_element_id,
                                title=text,
                                content=text,
                                section_path=section_path,
                                sort_order=sort_order,
                                content_hash=section_hash,
                                extracted_at=now,
                            )
                        )
                        heading_id_stack.append(section_id)
                        sort_order += 1
                        continue

                paragraph_hash = compute_content_hash(text)
                document_elements.append(
                    DocumentElementRow(
                        document_element_id=make_document_element_id(
                            source_file_id,
                            "paragraph",
                            None,
                            sort_order,
                            paragraph_hash,
                        ),
                        source_file_id=source_file_id,
                        element_type="paragraph",
                        parent_element_id=heading_id_stack[-1] if heading_id_stack else None,
                        content=text,
                        section_path="/".join(heading_stack) if heading_stack else None,
                        sort_order=sort_order,
                        content_hash=paragraph_hash,
                        extracted_at=now,
                    )
                )
                sort_order += 1
                continue

            if isinstance(block, Table):
                table_html = _render_table_html(block)
                table_text = " | ".join(
                    " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    for row in block.rows
                ).strip()
                table_hash = compute_content_hash(table_text)
                table_id = make_document_element_id(
                    source_file_id,
                    "table",
                    None,
                    sort_order,
                    table_hash,
                )
                current_section_path = "/".join(heading_stack) if heading_stack else None
                document_elements.append(
                    DocumentElementRow(
                        document_element_id=table_id,
                        source_file_id=source_file_id,
                        element_type="table",
                        parent_element_id=heading_id_stack[-1] if heading_id_stack else None,
                        title=path.stem,
                        content=table_text,
                        content_html=table_html,
                        section_path=current_section_path,
                        sort_order=sort_order,
                        content_hash=table_hash,
                        extracted_at=now,
                    )
                )
                sort_order += 1

                for row_index, row in enumerate(block.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()).strip()
                    if not row_text:
                        continue
                    row_html = "<tr>" + "".join(
                        f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells
                    ) + "</tr>"
                    row_hash = compute_content_hash(row_text)
                    document_elements.append(
                        DocumentElementRow(
                            document_element_id=make_document_element_id(
                                source_file_id,
                                "table_row",
                                None,
                                sort_order,
                                row_hash,
                            ),
                            source_file_id=source_file_id,
                            element_type="table_row",
                            parent_element_id=table_id,
                            content=row_text,
                            content_html=row_html,
                            section_path=current_section_path,
                            sort_order=sort_order,
                            row_index=row_index,
                            content_hash=row_hash,
                            extracted_at=now,
                        )
                    )
                    sort_order += 1

        return DocxExtractResult(
            source_file=source_file,
            document_elements=document_elements,
            hyperlinks=_extract_hyperlinks(document),
        )
