"""Unit tests for Sprint 2 extractors."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document

from fabric_kg_builder.model.ids import content_hash, make_document_element_id, make_source_file_id
from fabric_kg_builder.model.schemas import DocumentElementRow
from fabric_kg_builder.sources.chunker import Chunker
from fabric_kg_builder.sources.docx_extractor import DocxExtractor
from fabric_kg_builder.sources.html_extractor import HtmlExtractor
from fabric_kg_builder.sources.pdf_extractor import PdfExtractor
from fabric_kg_builder.sources.router import route
from fabric_kg_builder.sources.table_extractor import TableExtractor

MINIMAL_PDF = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>
stream
BT /F1 12 Tf 100 700 Td (Hello World) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000342 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
436
%%EOF"""

REAL_SURFACE_PDF = (
    Path(__file__).parent.parent.parent
    / "sample_data"
    / "Surface_Troubleshootings"
    / "Surface Pro 7 Kickstand Replacement Guide.pdf"
)


def _make_element(
    *,
    source_file_id: str,
    element_type: str,
    content: str,
    sort_order: int,
    content_html: str | None = None,
    parent_element_id: str | None = None,
    page_number: int | None = None,
    section_path: str | None = None,
) -> DocumentElementRow:
    value_hash = content_hash(content)
    return DocumentElementRow(
        document_element_id=make_document_element_id(
            source_file_id,
            element_type,
            page_number,
            sort_order,
            value_hash,
        ),
        source_file_id=source_file_id,
        element_type=element_type,
        parent_element_id=parent_element_id,
        content=content,
        content_html=content_html,
        page_number=page_number,
        section_path=section_path,
        sort_order=sort_order,
        content_hash=value_hash,
        extracted_at=datetime.now(timezone.utc),
    )


def test_pdf_extractor_with_minimal_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "tiny.pdf"
    pdf_path.write_bytes(MINIMAL_PDF)

    result = PdfExtractor.extract(pdf_path)

    assert result.page_count > 0
    assert result.document_elements
    assert any(element.element_type == "page" for element in result.document_elements)
    assert any("Hello World" in (element.content or "") for element in result.document_elements)


def test_docx_extractor(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph("This is a normal paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Qty"
    table.rows[1].cells[0].text = "Battery"
    table.rows[1].cells[1].text = "1"
    document.save(docx_path)

    result = DocxExtractor.extract(docx_path)

    section = next(element for element in result.document_elements if element.element_type == "section")
    paragraph = next(element for element in result.document_elements if element.element_type == "paragraph")
    table_element = next(element for element in result.document_elements if element.element_type == "table")

    assert section.title == "Overview"
    assert paragraph.content == "This is a normal paragraph."
    assert table_element.content_html is not None
    assert "<table>" in table_element.content_html


def test_html_extractor_from_string() -> None:
    html = """
    <html><body>
      <h1>Intro</h1>
      <p>Paragraph text.</p>
      <table><tr><th>Part</th><th>Qty</th></tr><tr><td>SSD</td><td>1</td></tr></table>
      <img src="surface.png" alt="Surface image" />
    </body></html>
    """

    result = HtmlExtractor.extract(html)

    types = [element.element_type for element in result.document_elements]
    assert "section" in types
    assert "paragraph" in types
    assert "table" in types
    assert "image_ref" in types
    table_element = next(element for element in result.document_elements if element.element_type == "table")
    image_element = next(element for element in result.document_elements if element.element_type == "image_ref")
    assert "<table>" in (table_element.content_html or "")
    assert "surface.png" in (image_element.content or "")


def test_chunker_mixed_elements() -> None:
    source_file_id = make_source_file_id("inline.html", content_hash("source"))
    table_id = make_document_element_id(source_file_id, "table", None, 2, content_hash("Cell 1 Cell 2"))
    elements = [
        _make_element(
            source_file_id=source_file_id,
            element_type="section",
            content="Overview",
            sort_order=0,
            section_path="Overview",
        ),
        _make_element(
            source_file_id=source_file_id,
            element_type="paragraph",
            content="NOTE: Replace battery carefully.",
            sort_order=1,
            section_path="Overview",
        ),
        _make_element(
            source_file_id=source_file_id,
            element_type="page",
            content="Raw page content.",
            sort_order=2,
            page_number=1,
        ),
        DocumentElementRow(
            document_element_id=table_id,
            source_file_id=source_file_id,
            element_type="table",
            content="Cell 1 Cell 2",
            content_html="<table><tr><td>Cell 1</td><td>Cell 2</td></tr></table>",
            sort_order=3,
            content_hash=content_hash("Cell 1 Cell 2"),
            extracted_at=datetime.now(timezone.utc),
        ),
        _make_element(
            source_file_id=source_file_id,
            element_type="table_row",
            content="Cell 1 | Cell 2",
            content_html="<tr><td>Cell 1</td><td>Cell 2</td></tr>",
            parent_element_id=table_id,
            sort_order=4,
        ),
    ]

    result = Chunker.extract(elements)
    chunk_types = [chunk.chunk_type for chunk in result.chunks]

    assert "section_text" in chunk_types
    assert "note" in chunk_types
    assert "raw_page_text" in chunk_types
    assert "table_html" in chunk_types
    assert "table_row" in chunk_types

    table_chunk = next(chunk for chunk in result.chunks if chunk.chunk_type == "table_html")
    note_chunk = next(chunk for chunk in result.chunks if chunk.chunk_type == "note")
    assert table_chunk.table_id == table_id
    assert table_chunk.embedding_text == "Cell 1 Cell 2"
    assert note_chunk.chunk_id == Chunker.extract(elements).chunks[1].chunk_id


def test_table_extractor_cells_from_table_element() -> None:
    source_file_id = make_source_file_id("inline.html", content_hash("source"))
    table_html = "<table><tr><th>Part</th><th>Qty</th></tr><tr><td>SSD</td><td>1</td></tr></table>"
    table_element = DocumentElementRow(
        document_element_id=make_document_element_id(
            source_file_id,
            "table",
            None,
            0,
            content_hash("Part Qty SSD 1"),
        ),
        source_file_id=source_file_id,
        element_type="table",
        content="Part Qty SSD 1",
        content_html=table_html,
        sort_order=0,
        content_hash=content_hash("Part Qty SSD 1"),
        extracted_at=datetime.now(timezone.utc),
    )

    cells = TableExtractor.extract([table_element])

    assert len(cells) == 4
    assert all(cell.element_type == "table_cell" for cell in cells)
    assert {(cell.row_index, cell.col_index) for cell in cells} == {(0, 0), (0, 1), (1, 0), (1, 1)}


def test_table_html_chunks() -> None:
    source_file_id = make_source_file_id("inline.html", content_hash("source"))
    table_elements = TableExtractor.extract_tables_from_html(
        "<table><tr><td>A</td><td>B</td></tr><tr><td>C</td><td>D</td></tr></table>",
        source_file_id,
    )

    chunks = TableExtractor.table_html_chunks(table_elements, source_file_id)

    assert any(chunk.chunk_type == "table_html" for chunk in chunks)
    assert any(chunk.chunk_type == "table_row" for chunk in chunks)
    table_chunk = next(chunk for chunk in chunks if chunk.chunk_type == "table_html")
    assert "<table>" in (table_chunk.content_html or "")


def test_router_extensions() -> None:
    assert route("guide.pdf") == "pdf_extractor"
    assert route("guide.docx") == "docx_extractor"
    assert route("guide.html") == "html_extractor"
    assert route("parts.csv") == "csv_loader"


@pytest.mark.integration
@pytest.mark.slow
def test_pdf_extractor_integration_real_surface_pdf() -> None:
    if not REAL_SURFACE_PDF.exists():
        pytest.skip(f"Fixture PDF not present: {REAL_SURFACE_PDF}")

    result = PdfExtractor.extract(REAL_SURFACE_PDF)

    assert len(result.document_elements) > 0
    assert result.page_count > 0
